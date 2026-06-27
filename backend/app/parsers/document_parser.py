"""
Resume and Job Description parsers.
Extracts raw text from PDF/DOCX, then structures it into sections.
"""
from __future__ import annotations
import re
import io
from typing import Optional

import fitz  # PyMuPDF
import docx  # python-docx
import pdfplumber


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text using PyMuPDF first; fall back to pdfplumber if empty."""
    text_parts = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
    except Exception:
        text_parts = []

    text = "\n".join(text_parts).strip()
    if text:
        return text

    # Fallback for tricky/scanned-ish PDFs with selectable text
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(p for p in parts if p.strip())


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {filename}")


# --- Section structuring -----------------------------------------------

SECTION_HEADERS = {
    "summary": ["summary", "objective", "profile"],
    "skills": ["skills", "technical skills", "core competencies"],
    "experience": ["experience", "work experience", "employment history", "professional experience"],
    "education": ["education", "academic background"],
    "projects": ["projects", "personal projects", "key projects"],
    "certifications": ["certifications", "certificates", "licenses"],
}


def _normalize_line(line: str) -> str:
    return re.sub(r"\s+", " ", line).strip()


def structure_resume(raw_text: str) -> dict:
    """
    Splits resume text into sections using common header keywords.
    This is heuristic, not perfect — good enough for scoring/keyword extraction.
    """
    lines = [_normalize_line(l) for l in raw_text.splitlines() if l.strip()]
    sections: dict[str, list[str]] = {key: [] for key in SECTION_HEADERS}
    sections["other"] = []

    current_section = "other"
    for line in lines:
        lowered = line.lower().strip(":").strip()
        matched = None
        for section, headers in SECTION_HEADERS.items():
            if lowered in headers or any(lowered == h for h in headers):
                matched = section
                break
        if matched:
            current_section = matched
            continue
        sections[current_section].append(line)

    name = lines[0] if lines else ""

    return {
        "name": name,
        "raw_text": raw_text,
        "summary": " ".join(sections["summary"]),
        "skills": sections["skills"],
        "experience": sections["experience"],
        "education": sections["education"],
        "projects": sections["projects"],
        "certifications": sections["certifications"],
        "other": sections["other"],
    }


def structure_job_description(raw_text: str) -> dict:
    """Lighter structuring for JDs: just keep raw text + naive section split."""
    lines = [_normalize_line(l) for l in raw_text.splitlines() if l.strip()]
    return {
        "raw_text": raw_text,
        "title_guess": lines[0] if lines else "",
        "lines": lines,
    }
